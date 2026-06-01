import streamlit as st
import tempfile
import os
import re
from transcriber import transcrever_audio
from exporter import gerar_docx, gerar_pdf, gerar_txt

st.set_page_config(page_title="Transc-Carac Pro", page_icon="📝", layout="wide")

# ==========================================
# 1. ESTILIZAÇÃO GLOBAL (CUSTOM CSS)
# ==========================================
def inject_custom_css():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700;800&display=swap');

    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }

    /* Ocultar menu do Streamlit e rodapé, mas manter o cabeçalho (para a seta da barra) */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {background-color: transparent !important;}
    
    /* Top padding reduction */
    .block-container {
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
        max-width: 1200px !important;
    }

    /* Estilo de Botões */
    .stButton>button {
        border-radius: 8px !important;
        transition: all 0.3s ease !important;
        font-weight: 600 !important;
        border: none !important;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06) !important;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgba(30, 78, 216, 0.2) !important;
        filter: brightness(1.05);
    }

    /* Cards para métricas */
    [data-testid="stMetric"] {
        background: #FFFFFF;
        border: 1px solid #E2E8F0;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        text-align: center;
        transition: all 0.3s ease;
    }
    [data-testid="stMetric"]:hover {
        transform: translateY(-5px);
        border-color: #1E4ED8;
        box-shadow: 0 10px 15px -3px rgba(30, 78, 216, 0.15);
    }
    [data-testid="stMetricLabel"] {
        font-size: 1rem !important;
        font-weight: 600 !important;
        color: #475569 !important;
        margin-bottom: 0.5rem !important;
        justify-content: center !important;
    }
    [data-testid="stMetricValue"] {
        font-size: 2.2rem !important;
        font-weight: 800 !important;
        background: -webkit-linear-gradient(45deg, #1E4ED8, #3B82F6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }

    /* Títulos e Headers */
    .main-title {
        background: linear-gradient(135deg, #1E4ED8, #3B82F6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3.5rem;
        font-weight: 800;
        margin-bottom: 0.5rem;
        line-height: 1.2;
    }
    .sub-title {
        color: #64748B;
        font-size: 1.2rem;
        font-weight: 400;
        margin-bottom: 2rem;
    }

    /* Steps */
    .step-title {
        color: #0F172A;
        font-weight: 700;
        font-size: 1.3rem;
        margin-top: 2rem;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        gap: 0.75rem;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #E2E8F0;
    }
    .step-icon {
        background: rgba(30, 78, 216, 0.1);
        color: #1E4ED8;
        width: 32px;
        height: 32px;
        display: flex;
        align-items: center;
        justify-content: center;
        border-radius: 8px;
        font-size: 1.1rem;
    }

    /* Tabs Customization */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
        background-color: transparent;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #FFFFFF;
        border-radius: 8px !important;
        gap: 0.5rem;
        padding: 0 1.5rem;
        font-weight: 600;
        color: #64748B;
        border: 1px solid #E2E8F0;
        transition: all 0.3s;
    }
    .stTabs [aria-selected="true"] {
        background-color: rgba(30, 78, 216, 0.05) !important;
        color: #1E4ED8 !important;
        border: 1px solid rgba(30, 78, 216, 0.3) !important;
    }
    
    /* Inputs, Selects e TextAreas (Forçando via Base Web) */
    div[data-baseweb="input"],
    div[data-baseweb="textarea"],
    div[data-baseweb="select"] {
        background-color: #FFFFFF !important;
        border: 2px solid #CBD5E1 !important;
        border-radius: 8px !important;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.02) !important;
        transition: all 0.3s ease !important;
    }
    
    div[data-baseweb="input"]:hover,
    div[data-baseweb="textarea"]:hover,
    div[data-baseweb="select"]:hover {
        border-color: #94A3B8 !important;
    }
    
    div[data-baseweb="input"]:focus-within,
    div[data-baseweb="textarea"]:focus-within,
    div[data-baseweb="select"]:focus-within {
        border-color: #1E4ED8 !important;
        box-shadow: 0 0 0 3px rgba(30, 78, 216, 0.15) !important;
    }
    
    /* Forçar a cor do texto nos campos internos */
    div[data-baseweb="input"] input,
    div[data-baseweb="textarea"] textarea {
        color: #0F172A !important;
        padding-top: 0.75rem !important;
        padding-bottom: 0.75rem !important;
        font-size: 1rem !important;
    }

    /* File uploader */
    [data-testid="stFileUploadDropzone"] {
        background-color: #FFFFFF !important;
        border: 2px dashed #94A3B8 !important;
        border-radius: 12px !important;
        padding: 2.5rem !important;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.02) !important;
        transition: all 0.3s ease !important;
    }
    [data-testid="stFileUploadDropzone"]:hover {
        border-color: #1E4ED8 !important;
        background-color: rgba(30, 78, 216, 0.02) !important;
        box-shadow: 0 8px 16px rgba(30, 78, 216, 0.08) !important;
    }
    /* Reduzir o espaço em branco no topo da barra lateral para subir a logo */
    [data-testid="stSidebarHeader"] {
        padding: 0 !important;
        min-height: 0 !important;
    }
    [data-testid="stSidebarUserContent"] {
        padding-top: 0rem !important;
    }
    [data-testid="stSidebar"] [data-testid="stImage"] {
        margin-top: -2rem !important;
    }
    </style>
    """, unsafe_allow_html=True)

inject_custom_css()

# Cabeçalho Principal
st.markdown("<div class='main-title'>Transc-Carac Pro</div>", unsafe_allow_html=True)
st.markdown("<div class='sub-title'>Sistema inteligente alimentado por IA para <b>transcrição de áudio</b> e <b>análise profunda de texto</b>.</div>", unsafe_allow_html=True)
st.markdown("<br>", unsafe_allow_html=True)

tabs = st.tabs(["🎙️ Transcrição de Áudio", "📊 Contagem de Caracteres"])

# ==========================================
# 2. BARRA LATERAL (SIDEBAR) MAIS LIMPA
# ==========================================
with st.sidebar:
    try:
        # Usa a logo da CM Group (caminho relativo a onde o streamlit foi executado)
        st.image("00 - Logo da CM/Logo C&M Group.png", use_container_width=True)
    except Exception:
        # Fallback caso rode de outra pasta
        st.markdown("<h1 style='text-align: center; color: #1E4ED8;'>🎙️</h1>", unsafe_allow_html=True)
    
    st.header("⚙️ Opções Principais")
    idioma = st.text_input("🌍 Idioma do Áudio (ex: pt, en)", value="pt")
    
    st.markdown("---")
    
    # Escondendo as configurações sensíveis em um expander
    with st.expander("🛠️ Configurações Avançadas"):
        st.markdown("**Diarização (Separação de Vozes)**")
        hf_token = st.text_input("🔑 HuggingFace Token", type="password", help="Cole o token do HuggingFace para separar múltiplos locutores automaticamente.")
        num_locutores = st.number_input("👥 Número de Locutores (0 = auto)", min_value=0, value=0, step=1)
        
        st.markdown("**Qualidade da Transcrição**")
        modelo_whisper = st.selectbox("Qualidade/Modelo", ["tiny", "base", "small", "medium", "large"], index=1, help="Maior qualidade = tempo de processamento mais longo.")

# ==========================================
# 3. ABA 1: TRANSCRIÇÃO PASSO A PASSO
# ==========================================
with tabs[0]:
    
    st.markdown("<div class='step-title'><div class='step-icon'>1</div> Selecione o Áudio</div>", unsafe_allow_html=True)
    uploaded_audio = st.file_uploader("Faça upload do arquivo de áudio ou vídeo (MP3, WAV, MP4)", type=['mp3', 'wav', 'mp4', 'm4a', 'ogg'], label_visibility="collapsed")
    
    if uploaded_audio:
        # Layout limpo para o player
        col_player, _ = st.columns([1, 1])
        with col_player:
            st.audio(uploaded_audio)
            
        # Limpar estado se arquivo mudar
        if 'ultimo_arquivo' not in st.session_state or st.session_state['ultimo_arquivo'] != uploaded_audio.name:
            st.session_state['ultimo_arquivo'] = uploaded_audio.name
            if 'transcricao' in st.session_state:
                del st.session_state['transcricao']
                
        st.markdown("<div class='step-title'><div class='step-icon'>2</div> Iniciar Transcrição</div>", unsafe_allow_html=True)
        col_btn, _ = st.columns([1, 2])
        with col_btn:
            iniciar = st.button("🚀 Iniciar Transcrição Agora", type="primary", use_container_width=True)
                
        if iniciar:
            with st.spinner("Processando o áudio com Inteligência Artificial..."):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                def update_progress(percent, msg):
                    progress_bar.progress(percent)
                    status_text.text(msg)
                
                # Salvar temporariamente
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_audio.name)[1]) as tmp_file:
                    tmp_file.write(uploaded_audio.getbuffer())
                    temp_audio_path = tmp_file.name
                
                try:
                    resultado = transcrever_audio(
                        temp_audio_path,
                        modelo=modelo_whisper,
                        idioma=idioma,
                        hf_token=hf_token if hf_token.strip() else None,
                        num_locutores=num_locutores if num_locutores > 0 else None,
                        progress_callback=update_progress
                    )
                    st.session_state['transcricao'] = resultado
                    st.success("✅ Transcrição concluída com sucesso!")
                except Exception as e:
                    st.error(f"Erro durante a transcrição: {e}")
                finally:
                    if os.path.exists(temp_audio_path):
                        os.remove(temp_audio_path)
                        
    # ==========================================
    # 4. RESULTADOS E EXPORTAÇÃO
    # ==========================================
    if 'transcricao' in st.session_state:
        st.markdown("<div class='step-title'><div class='step-icon'>3</div> Resultados e Exportação</div>", unsafe_allow_html=True)
        
        texto_t = st.session_state['transcricao']
        st.text_area("Texto gerado:", value=texto_t, height=300, label_visibility="collapsed")
        
        # Estatísticas via st.metric
        qtd_com_espaco = len(texto_t)
        qtd_sem_espaco = len(texto_t.replace(" ", "").replace("\n", "").replace("\t", ""))
        palavras = len(re.findall(r'\b\w+\b', texto_t))
        
        st.markdown("### 📊 Estatísticas da Transcrição")
        m1, m2, m3 = st.columns(3)
        m1.metric("Caracteres (Com Espaços)", f"{qtd_com_espaco:,}".replace(",", "."))
        m2.metric("Caracteres (Sem Espaços)", f"{qtd_sem_espaco:,}".replace(",", "."))
        m3.metric("Total de Palavras", f"{palavras:,}".replace(",", "."))
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        # Botões de Exportação Lado a Lado
        ex1, ex2, ex3 = st.columns(3)
        with ex1:
            docx_buffer = gerar_docx(texto_t)
            st.download_button("📄 Baixar em Word (.docx)", data=docx_buffer, file_name="transcricao.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="secondary")
        with ex2:
            pdf_buffer = gerar_pdf(texto_t)
            st.download_button("📕 Baixar em PDF (.pdf)", data=pdf_buffer, file_name="transcricao.pdf", mime="application/pdf", use_container_width=True, type="secondary")
        with ex3:
            txt_buffer = gerar_txt(texto_t)
            st.download_button("📝 Baixar em Texto (.txt)", data=txt_buffer, file_name="transcricao.txt", mime="text/plain", use_container_width=True, type="primary")

# ==========================================
# 5. ABA 2: CONTAGEM DE CARACTERES
# ==========================================
with tabs[1]:
    st.markdown("<div class='step-title'><div class='step-icon'>📊</div> Contagem de Caracteres e Palavras</div>", unsafe_allow_html=True)
    st.markdown("<div style='color: #64748B; margin-bottom: 1rem;'>Cole o seu texto abaixo ou faça upload de um arquivo para análise instantânea.</div>", unsafe_allow_html=True)
    
    texto_para_contar = ""
    
    col_input1, col_input2 = st.columns([1.5, 1])
    with col_input1:
        texto_digitado = st.text_area("Área de Texto:", height=250, placeholder="Cole ou digite seu texto aqui...")
    with col_input2:
        st.info("Ou use um arquivo:")
        arquivo_texto = st.file_uploader("Upload de arquivo (txt, pdf, docx, doc)", type=['txt', 'pdf', 'docx', 'doc'], label_visibility="collapsed")
    
    if arquivo_texto:
        nome_arquivo = arquivo_texto.name.lower()
        try:
            if nome_arquivo.endswith('.txt'):
                texto_para_contar = arquivo_texto.getvalue().decode("utf-8", errors="ignore")
            elif nome_arquivo.endswith('.pdf'):
                import pypdf
                leitor = pypdf.PdfReader(arquivo_texto)
                texto_extraido = []
                for pagina in leitor.pages:
                    texto_pagina = pagina.extract_text()
                    if texto_pagina:
                        texto_extraido.append(texto_pagina)
                texto_para_contar = "\n".join(texto_extraido)
            elif nome_arquivo.endswith('.docx'):
                import docx
                import io
                doc = docx.Document(io.BytesIO(arquivo_texto.getvalue()))
                texto_para_contar = "\n".join([p.text for p in doc.paragraphs])
            elif nome_arquivo.endswith('.doc'):
                import subprocess
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
                    tmp.write(arquivo_texto.getvalue())
                    tmp_path = tmp.name
                try:
                    resultado = subprocess.run(['antiword', tmp_path], capture_output=True, text=True)
                    if resultado.returncode == 0:
                        texto_para_contar = resultado.stdout
                    else:
                        st.error("Erro ao ler arquivo .doc. Verifique se o formato está correto.")
                        texto_para_contar = ""
                except FileNotFoundError:
                    try:
                        import docx
                        import io
                        doc = docx.Document(io.BytesIO(arquivo_texto.getvalue()))
                        texto_para_contar = "\n".join([p.text for p in doc.paragraphs])
                    except:
                        st.error("Pacote 'antiword' não encontrado para ler .doc nativamente.")
                        texto_para_contar = ""
                finally:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
            else:
                st.error("Formato de arquivo não suportado.")
        except Exception as e:
            st.error(f"Erro ao processar o arquivo: {e}")
            texto_para_contar = ""
            
    elif texto_digitado:
        texto_para_contar = texto_digitado
        
    if texto_para_contar:
        qtd_com_espaco = len(texto_para_contar)
        qtd_sem_espaco = len(texto_para_contar.replace(" ", "").replace("\n", "").replace("\t", ""))
        palavras = len(re.findall(r'\b\w+\b', texto_para_contar))
        
        st.markdown("---")
        st.markdown("### Resultados da Análise")
        r1, r2, r3 = st.columns(3)
        r1.metric("Caracteres (Com Espaços)", f"{qtd_com_espaco:,}".replace(",", "."))
        r2.metric("Caracteres (Sem Espaços)", f"{qtd_sem_espaco:,}".replace(",", "."))
        r3.metric("Total de Palavras", f"{palavras:,}".replace(",", "."))
