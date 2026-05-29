import io
from docx import Document
from fpdf import FPDF

def gerar_docx(texto_transcrito):
    '''Gera um documento DOCX em memória a partir do texto transcrito.'''
    doc = Document()
    doc.add_heading("Transcrição de Áudio", 0)
    
    # Adiciona os parágrafos
    for linha in texto_transcrito.split('\n'):
        if linha.strip():
            doc.add_paragraph(linha.strip())
            
    # Salva no buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_txt(texto_transcrito):
    '''Gera um arquivo TXT em memória a partir do texto transcrito.'''
    buffer = io.BytesIO()
    buffer.write(texto_transcrito.encode('utf-8'))
    buffer.seek(0)
    return buffer

class PDF(FPDF):
    def header(self):
        self.set_font("helvetica", "B", 15)
        self.cell(0, 10, "Transcrição de Áudio", align="C", ln=True)
        self.ln(10)

def gerar_pdf(texto_transcrito):
    '''Gera um documento PDF em memória a partir do texto transcrito.'''
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    pdf.set_font("helvetica", size=12)
    
    # O texto precisa ser codificado adequadamente para a fonte base
    for linha in texto_transcrito.split('\n'):
        if linha.strip():
            # Limpar caracteres que podem quebrar o fpdf com fontes padrão
            txt_safe = linha.strip().encode('latin-1', 'replace').decode('latin-1')
            # multi_cell lida com quebra de linha
            pdf.multi_cell(0, 10, txt=txt_safe)
            pdf.ln(2)
            
    buffer = io.BytesIO()
    # fpdf2 permite output para bytes
    pdf_bytes = pdf.output(dest='S')
    buffer.write(pdf_bytes)
    buffer.seek(0)
    return buffer
